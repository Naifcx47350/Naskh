import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from app.config import get_settings
from app.schemas import DocumentExtraction

_ARABIC_RE = re.compile(r"[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]")


def _is_arabic(text: str) -> bool:
    return bool(_ARABIC_RE.search(text or ""))


def _set_paragraph_rtl(paragraph: Paragraph) -> None:
    """Mark a paragraph (and its runs) right-to-left.

    Word performs its own Arabic shaping and bidi reordering, so the text is
    stored logically and only flagged as RTL. Pre-shaping with arabic-reshaper
    / python-bidi here would double-process the text and render it broken in
    Word; those libraries belong in image/PDF rendering paths instead.
    """
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")

    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for run in paragraph.runs:
        r_pr = run._r.get_or_add_rPr()
        rtl = r_pr.find(qn("w:rtl"))
        if rtl is None:
            rtl = OxmlElement("w:rtl")
            r_pr.append(rtl)
        rtl.set(qn("w:val"), "1")


def _add_paragraph(doc: Document, text: str, style: str | None = None) -> Paragraph:
    paragraph = doc.add_paragraph(text, style=style) if style else doc.add_paragraph(text)
    if _is_arabic(text):
        _set_paragraph_rtl(paragraph)
    return paragraph


def _set_cell_text(cell, text: str) -> None:
    cell.text = text
    if _is_arabic(text):
        for paragraph in cell.paragraphs:
            _set_paragraph_rtl(paragraph)


def write_docx(document_id: str, extraction: DocumentExtraction) -> Path:
    output_path = get_settings().data_dir / "exports" / f"{document_id}.docx"
    doc = Document()
    doc.add_heading("Naskh Digitized Document", level=1)
    _add_paragraph(doc, extraction.summary)

    if extraction.transcription:
        doc.add_heading("Editable Transcription", level=2)
        _add_paragraph(doc, extraction.transcription)

    if extraction.fields:
        doc.add_heading("Structured Fields", level=2)
        table = doc.add_table(rows=1, cols=4)
        header = table.rows[0].cells
        header[0].text = "Field"
        header[1].text = "Value"
        header[2].text = "Source"
        header[3].text = "Confidence"

        for field in extraction.fields:
            row = table.add_row().cells
            _set_cell_text(row[0], field.label)
            _set_cell_text(row[1], field.value)
            _set_cell_text(row[2], field.source.snippet)
            _set_cell_text(row[3], f"{field.confidence:.0%}")

    if extraction.notes:
        doc.add_heading("Human Review Notes", level=2)
        for note in extraction.notes:
            _add_paragraph(doc, note, style="List Bullet")

    doc.save(output_path)
    return output_path


def write_json(document_id: str, extraction: DocumentExtraction) -> Path:
    output_path = get_settings().data_dir / "exports" / f"{document_id}.json"
    output_path.write_text(
        json.dumps(extraction.model_dump(), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return output_path
